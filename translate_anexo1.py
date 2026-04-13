from docx import Document

SRC = "ASAS DE SOCORRO - ANEXO I - 24.03.26 (1) (1).docx"
DST = "ASAS DE SOCORRO - ANNEX I - 24.03.26 (EN).docx"

TRANSLATIONS = [
    ("ANEXO N.\u00ba 1", "ANNEX No. 1"),
    ("AO ACORDO DE COOPERA\u00c7\u00c3O INTERNACIONAL", "TO THE INTERNATIONAL COOPERATION AGREEMENT"),
    ("Anexo n.\u00ba 1 ao Acordo de Coopera\u00e7\u00e3o Internacional celebrado entre ASSOCIA\u00c7\u00c3O ASAS DE SOCORRO e AVIATION DEVELOPERS, INC., disciplinando a cust\u00f3dia tempor\u00e1ria da aeronave Cessna T210G, matr\u00edcula FAA N230M, com interven\u00e7\u00e3o anuente da MISSION AVIATION FELLOWSHIP.",
     "Annex No. 1 to the International Cooperation Agreement entered into between ASSOCIA\u00c7\u00c3O ASAS DE SOCORRO and AVIATION DEVELOPERS, INC., governing the temporary custody of the Cessna T210G aircraft, FAA registration N230M, with the consenting intervention of MISSION AVIATION FELLOWSHIP."),
    ("CONSIDERANDO que a aeronave Cessna T210G, matr\u00edcula FAA N230M, n\u00famero de s\u00e9rie U20600634, ano de fabrica\u00e7\u00e3o 1964, foi doada pela MISSION AVIATION FELLOWSHIP (\u201cMAF\u201d) \u00e0 ASSOCIA\u00c7\u00c3O ASAS DE SOCORRO, que desde ent\u00e3o \u00e9 sua propriet\u00e1ria, n\u00e3o obstante a titularidade registral perante a FAA permanecer, formalmente, em nome da MAF;",
     "WHEREAS the Cessna T210G aircraft, FAA registration N230M, serial number U20600634, year of manufacture 1964, was donated by MISSION AVIATION FELLOWSHIP (\u201cMAF\u201d) to ASSOCIA\u00c7\u00c3O ASAS DE SOCORRO, which has since been its owner, notwithstanding that legal title before the FAA formally remains in the name of MAF;"),
    ("CONSIDERANDO que, por raz\u00f5es burocr\u00e1ticas e operacionais relacionadas ao procedimento de registro perante a FAA, a transfer\u00eancia do t\u00edtulo legal da aeronave N230M ser\u00e1 realizada diretamente da MAF para a AVIATION DEVELOPERS, INC. (\u201cADI\u201d), para que esta a custodie temporariamente e, ao final, transfira o t\u00edtulo legal \u00e0 ASAS DE SOCORRO ou a entidade por ela designada;",
     "WHEREAS, for bureaucratic and operational reasons related to the FAA registration procedure, the transfer of legal title to aircraft N230M will be made directly from MAF to AVIATION DEVELOPERS, INC. (\u201cADI\u201d), so that ADI may temporarily hold custody and, at the end, transfer legal title to ASAS DE SOCORRO or an entity designated by it;"),
    ("CONSIDERANDO que a ADI deter\u00e1 o t\u00edtulo legal da aeronave N230M exclusivamente na condi\u00e7\u00e3o de fiel deposit\u00e1ria e custodiante tempor\u00e1ria, sem que tal circunst\u00e2ncia implique qualquer direito de propriedade, frui\u00e7\u00e3o ou disposi\u00e7\u00e3o sobre a aeronave;",
     "WHEREAS ADI will hold legal title to aircraft N230M exclusively as a fiduciary and temporary custodian, such circumstance not implying any right of ownership, use, or disposition over the aircraft;"),
    ("CONSIDERANDO que a aeronave N230M encontra-se em fase final de prepara\u00e7\u00e3o para voo nos Estados Unidos da Am\u00e9rica, com previs\u00e3o de transfer\u00eancia definitiva ao Brasil ap\u00f3s a conclus\u00e3o das etapas de finaliza\u00e7\u00e3o descritas neste Anexo;",
     "WHEREAS aircraft N230M is in the final stages of flight preparation in the United States of America, with definitive transfer to Brazil expected after completion of the finalization steps described in this Annex;"),
    ("Resolvem as PARCEIRAS e a MAF celebrar o presente ANEXO N.\u00ba 1, que se reger\u00e1 pelas cl\u00e1usulas e condi\u00e7\u00f5es a seguir:",
     "The PARTNERS and MAF hereby agree to enter into this ANNEX No. 1, which shall be governed by the following clauses and conditions:"),
    ("CL\u00c1USULA PRIMEIRA \u2013 DAS PARTES", "CLAUSE ONE \u2013 THE PARTIES"),
    ("1.1. S\u00e3o partes do presente Anexo:", "1.1. The parties to this Annex are:"),
    ("(a) ASSOCIA\u00c7\u00c3O ASAS DE SOCORRO, doravante denominada ASAS DE SOCORRO, associa\u00e7\u00e3o civil sem fins lucrativos, inscrita no CNPJ/MF sob o n.\u00ba 01.052.752/0001-69, com sede na Avenida Presidente Juscelino Kubistchek, qd. 08, Lt. 13, Setor Aeroporto, An\u00e1polis/GO, CEP 75.104-280, neste ato representada na forma de seu Estatuto Social;",
     "(a) ASSOCIA\u00c7\u00c3O ASAS DE SOCORRO, hereinafter referred to as ASAS DE SOCORRO, a non-profit civil association, registered under CNPJ/MF No. 01.052.752/0001-69, headquartered at Avenida Presidente Juscelino Kubistchek, block 08, lot 13, Setor Aeroporto, An\u00e1polis/GO, CEP 75.104-280, hereby represented in accordance with its Articles of Association;"),
    ("(b) AVIATION DEVELOPERS, INC., doravante denominada ADI, corpora\u00e7\u00e3o sem fins lucrativos constitu\u00edda nos termos da Se\u00e7\u00e3o 501(c)(3) do Internal Revenue Code dos Estados Unidos da Am\u00e9rica, registrada no Estado de Missouri, com sede em 1730 State Highway HH, Highlandville, MO 65669, EUA, neste ato representada na forma de seus Bylaws; e",
     "(b) AVIATION DEVELOPERS, INC., hereinafter referred to as ADI, a non-profit corporation organized under Section 501(c)(3) of the Internal Revenue Code of the United States of America, registered in the State of Missouri, with its principal office at 1730 State Highway HH, Highlandville, MO 65669, USA, hereby represented in accordance with its Bylaws; and"),
    ("(c) MISSION AVIATION FELLOWSHIP, doravante denominada MAF, corpora\u00e7\u00e3o sem fins lucrativos sediada em 112 N Pilatus Ln, Nampa, Idaho 83687, EUA, neste ato representada na forma de seus documentos constitutivos, na qualidade de interveniente-anuente.",
     "(c) MISSION AVIATION FELLOWSHIP, hereinafter referred to as MAF, a non-profit corporation headquartered at 112 N Pilatus Ln, Nampa, Idaho 83687, USA, hereby represented in accordance with its governing documents, in its capacity as consenting intervening party."),
    ("CL\u00c1USULA SEGUNDA \u2013 DA AERONAVE", "CLAUSE TWO \u2013 THE AIRCRAFT"),
    ("2.1. O presente Anexo tem por objeto a cust\u00f3dia tempor\u00e1ria, pela ADI, da seguinte aeronave:",
     "2.1. The subject matter of this Annex is the temporary custody by ADI of the following aircraft:"),
    ("Fabricante / Modelo", "Manufacturer / Model"),
    ("Matr\u00edcula FAA", "FAA Registration"),
    ("N\u00famero de S\u00e9rie", "Serial Number"),
    ("Ano de Fabrica\u00e7\u00e3o", "Year of Manufacture"),
    ("Origem do t\u00edtulo legal atual", "Current Legal Title Holder"),
    ("Propriet\u00e1ria", "Owner"),
    ("2.2. ASAS DE SOCORRO \u00e9 e permanece propriet\u00e1ria da aeronave N230M. A deten\u00e7\u00e3o do t\u00edtulo legal pela ADI \u00e9 de car\u00e1ter estritamente tempor\u00e1rio e instrumental, para os fins previstos neste Anexo, n\u00e3o implicando qualquer direito de propriedade, frui\u00e7\u00e3o ou disposi\u00e7\u00e3o em favor da ADI.",
     "2.2. ASAS DE SOCORRO is and remains the owner of aircraft N230M. ADI\u2019s holding of legal title is strictly temporary and instrumental in nature, for the purposes set forth in this Annex, and does not confer any right of ownership, use, or disposition upon ADI."),
    ("CL\u00c1USULA TERCEIRA \u2013 DA NU\u00caNCA DA MAF", "CLAUSE THREE \u2013 MAF\u2019S CONSENT"),
    ("CL\u00c1USULA TERCEIRA \u2013 DA NU\u00caNCIA DA MAF", "CLAUSE THREE \u2013 MAF\u2019S CONSENT"),
    ("CL\u00c1USULA TERCEIRA \u2013 DA NU\u00caN\u00c7A DA MAF", "CLAUSE THREE \u2013 MAF\u2019S CONSENT"),
    ("3.1. MAF, na qualidade de atual detentora do t\u00edtulo legal da aeronave N230M perante a FAA, declara, para todos os fins de direito, que: (a) reconhece que ASAS DE SOCORRO \u00e9 a propriet\u00e1ria da aeronave N230M desde a respectiva doa\u00e7\u00e3o, n\u00e3o possuindo a MAF qualquer direito real, obriga\u00e7\u00e3o ou responsabilidade sobre a aeronave al\u00e9m da titularidade registral perante a FAA; (b) consente expressamente com a transfer\u00eancia do t\u00edtulo legal da aeronave N230M da MAF para a ADI, a ser formalizada perante a FAA; e (c) declara n\u00e3o existir \u00f4nus, gravames, penhoras, hipotecas ou quaisquer restri\u00e7\u00f5es sobre a aeronave N230M que possam impedir ou onerar a transfer\u00eancia ora consentida.",
     "3.1. MAF, as the current holder of legal title to aircraft N230M before the FAA, hereby declares, for all legal purposes, that: (a) it acknowledges that ASAS DE SOCORRO has been the owner of aircraft N230M since its donation, and that MAF holds no real right, obligation, or liability with respect to the aircraft beyond its registered title before the FAA; (b) it expressly consents to the transfer of legal title to aircraft N230M from MAF to ADI, to be formalized before the FAA; and (c) it declares that there are no liens, encumbrances, attachments, mortgages, or any other restrictions on aircraft N230M that could prevent or burden the transfer hereby consented to."),
    ("3.2. A assinatura da MAF neste instrumento n\u00e3o lhe imputa qualquer obriga\u00e7\u00e3o futura de natureza operacional, financeira ou regulat\u00f3ria relacionada \u00e0 aeronave N230M.",
     "3.2. MAF\u2019s signature on this instrument does not impose upon it any future obligation of an operational, financial, or regulatory nature related to aircraft N230M."),
    ("CL\u00c1USULA QUARTA \u2013 DA FINALIDADE DA CUST\u00d3DIA", "CLAUSE FOUR \u2013 PURPOSE OF CUSTODY"),
    ("4.1. ADI manter\u00e1 o t\u00edtulo legal da aeronave N230M exclusivamente para os seguintes fins: (a) conclus\u00e3o dos servi\u00e7os de montagem, revis\u00e3o e prepara\u00e7\u00e3o para voo; (b) instala\u00e7\u00e3o de r\u00e1dios, avi\u00f4nicos e instrumentos de voo; (c) modifica\u00e7\u00e3o de assentos conforme padr\u00e3o operacional de ASAS DE SOCORRO; (d) capta\u00e7\u00e3o de recursos financeiros nos Estados Unidos para cobertura dos custos de finaliza\u00e7\u00e3o da aeronave, a serem repassados integralmente a ASAS DE SOCORRO nos termos da Cl\u00e1usula Sexta; e (e) prepara\u00e7\u00e3o para exporta\u00e7\u00e3o e registro definitivo no Brasil perante a ANAC.",
     "4.1. ADI shall hold legal title to aircraft N230M exclusively for the following purposes: (a) completion of assembly, overhaul, and flight preparation services; (b) installation of radios, avionics, and flight instruments; (c) seat modification in accordance with ASAS DE SOCORRO\u2019s operational standards; (d) fundraising in the United States to cover the costs of finalizing the aircraft, which shall be remitted in full to ASAS DE SOCORRO pursuant to Clause Six; and (e) preparation for export and definitive registration in Brazil before ANAC."),
    ("4.2. \u00c9 vedado \u00e0 ADI: (a) alienar, ceder, transferir ou de qualquer forma dispor da aeronave a terceiros; (b) onerar, hipotecar, penhorar ou oferecer a aeronave como garantia; (c) utiliz\u00e1-la para finalidade diversa das previstas neste Anexo; e (d) permitir a opera\u00e7\u00e3o por terceiros n\u00e3o autorizados por ASAS DE SOCORRO.",
     "4.2. ADI is prohibited from: (a) selling, assigning, transferring, or otherwise disposing of the aircraft to third parties; (b) encumbering, mortgaging, pledging, or offering the aircraft as collateral; (c) using it for any purpose other than those set forth in this Annex; and (d) allowing operation by third parties not authorized by ASAS DE SOCORRO."),
    ("CL\u00c1USULA QUINTA \u2013 DO CRONOGRAMA OPERACIONAL", "CLAUSE FIVE \u2013 OPERATIONAL SCHEDULE"),
    ("5.1. O itiner\u00e1rio previsto para a aeronave N230M \u00e9 o seguinte: (a) Minnesota: hangar atual, onde a aeronave se encontra para conclus\u00e3o da montagem e prepara\u00e7\u00e3o para voo; (b) Missouri (base da ADI): instala\u00e7\u00e3o de r\u00e1dios, avi\u00f4nicos e instrumentos de voo; (c) JAARS: modifica\u00e7\u00e3o de assentos conforme padr\u00e3o de ASAS DE SOCORRO; e (d) Brasil: transfer\u00eancia definitiva, ap\u00f3s exporta\u00e7\u00e3o e registro perante a ANAC.",
     "5.1. The planned itinerary for aircraft N230M is as follows: (a) Minnesota: current hangar, where the aircraft is located for completion of assembly and flight preparation; (b) Missouri (ADI base): installation of radios, avionics, and flight instruments; (c) JAARS: seat modification in accordance with ASAS DE SOCORRO\u2019s standards; and (d) Brazil: definitive transfer, after export and registration before ANAC."),
    ("5.2. ADI comunicar\u00e1 a ASAS DE SOCORRO, com anteced\u00eancia m\u00ednima de 5 (cinco) dias \u00fateis, qualquer altera\u00e7\u00e3o no itiner\u00e1rio ou nos prazos previstos nesta cl\u00e1usula.",
     "5.2. ADI shall notify ASAS DE SOCORRO at least 5 (five) business days in advance of any change to the itinerary or deadlines set forth in this clause."),
    ("CL\u00c1USULA SEXTA \u2013 DA CAPTA\u00c7\u00c3O DE RECURSOS", "CLAUSE SIX \u2013 FUNDRAISING"),
    ("6.1. ADI conduzir\u00e1 atividades de capta\u00e7\u00e3o de recursos (\u201cfundraising\u201d) nos Estados Unidos para cobertura dos custos de finaliza\u00e7\u00e3o da aeronave N230M.",
     "6.1. ADI shall conduct fundraising activities in the United States to cover the costs of finalizing aircraft N230M."),
    ("6.2. Os recursos captados ser\u00e3o repassados integralmente a ASAS DE SOCORRO, sendo vedada qualquer dedu\u00e7\u00e3o de taxa administrativa (\u201coverhead\u201d), nos termos da Cl\u00e1usula 3\u00aa do Acordo de Coopera\u00e7\u00e3o Internacional.",
     "6.2. All funds raised shall be remitted in full to ASAS DE SOCORRO, with no administrative fee deduction permitted, pursuant to Clause 3 of the International Cooperation Agreement."),
    ("6.3. ADI encaminhar\u00e1 a ASAS DE SOCORRO relat\u00f3rio discriminando os valores captados, as dedu\u00e7\u00f5es autorizadas e os repasses efetuados, nos termos da Cl\u00e1usula 3.2 do Acordo de Coopera\u00e7\u00e3o Internacional.",
     "6.3. ADI shall submit to ASAS DE SOCORRO a report itemizing the amounts raised, authorized deductions, and remittances made, pursuant to Clause 3.2 of the International Cooperation Agreement."),
    ("CL\u00c1USULA S\u00c9TIMA \u2013 DA OPERA\u00c7\u00c3O E DO SEGURO", "CLAUSE SEVEN \u2013 OPERATION AND INSURANCE"),
    ("7.1. Somente poder\u00e1 operar a aeronave N230M piloto indicado ou designado por ASAS DE SOCORRO.",
     "7.1. Only a pilot indicated or designated by ASAS DE SOCORRO may operate aircraft N230M."),
    ("7.2. ASAS DE SOCORRO poder\u00e1, a seu exclusivo crit\u00e9rio, contratar e custear seguro de voo para a aeronave N230M. Havendo ap\u00f3lice vigente, ADI obriga-se a cumprir todas as suas condi\u00e7\u00f5es e, se necess\u00e1rio, a atender os requisitos exigidos pela seguradora que lhe sejam imput\u00e1veis.",
     "7.2. ASAS DE SOCORRO may, at its sole discretion, obtain and pay for flight insurance for aircraft N230M. If a policy is in force, ADI undertakes to comply with all its terms and, if required, to meet any insurer requirements attributable to it."),
    ("CL\u00c1USULA OITAVA \u2013 DA DEVOLU\u00c7\u00c3O DO T\u00cdTULO LEGAL", "CLAUSE EIGHT \u2013 RETURN OF LEGAL TITLE"),
    ("8.1. Conclu\u00eddas as atividades previstas na Cl\u00e1usula Quarta, ou na hip\u00f3tese de extin\u00e7\u00e3o antecipada deste Anexo, ADI executar\u00e1 prontamente todos os documentos de transfer\u00eancia do t\u00edtulo legal da aeronave N230M a ASAS DE SOCORRO ou entidade por ela designada.",
     "8.1. Upon completion of the activities set forth in Clause Four, or in the event of early termination of this Annex, ADI shall promptly execute all documents required to transfer legal title to aircraft N230M to ASAS DE SOCORRO or an entity designated by it."),
    ("8.2. A transfer\u00eancia dever\u00e1 ser conclu\u00edda no prazo m\u00e1ximo de 60 (sessenta) dias contados da conclus\u00e3o das atividades ou da extin\u00e7\u00e3o do Anexo, o que ocorrer primeiro.",
     "8.2. The transfer shall be completed within a maximum of 60 (sixty) days from the completion of activities or termination of the Annex, whichever occurs first."),
    ("8.3. ADI indenizar\u00e1 ASAS DE SOCORRO de quaisquer perdas, danos ou responsabilidades decorrentes de atos ou omiss\u00f5es durante a cust\u00f3dia, incluindo eventual descumprimento das veda\u00e7\u00f5es previstas no item 4.2.",
     "8.3. ADI shall indemnify ASAS DE SOCORRO for any losses, damages, or liabilities arising from acts or omissions during the custody period, including any breach of the prohibitions set forth in item 4.2."),
    ("CL\u00c1USULA NONA \u2013 DA VIG\u00caNCIA", "CLAUSE NINE \u2013 TERM"),
    ("9.1. O presente Anexo entra em vigor na data de sua assinatura e vigorar\u00e1 at\u00e9 a conclus\u00e3o de todas as obriga\u00e7\u00f5es nele previstas e a transfer\u00eancia definitiva do t\u00edtulo legal da aeronave \u00e0 ASAS DE SOCORRO.",
     "9.1. This Annex enters into force on the date of its signature and shall remain in effect until the completion of all obligations set forth herein and the definitive transfer of legal title to the aircraft to ASAS DE SOCORRO."),
    ("9.2. O presente Anexo poder\u00e1 ser rescindido de pleno direito por ASAS DE SOCORRO, independentemente de notifica\u00e7\u00e3o pr\u00e9via, nas hip\u00f3teses previstas na Cl\u00e1usula 11.3 do Acordo de Coopera\u00e7\u00e3o Internacional.",
     "9.2. This Annex may be terminated by operation of law by ASAS DE SOCORRO, without prior notice, in the circumstances set forth in Clause 11.3 of the International Cooperation Agreement."),
    ("CL\u00c1USULA D\u00c9CIMA \u2013 DAS DISPOSI\u00c7\u00d5ES FINAIS", "CLAUSE TEN \u2013 FINAL PROVISIONS"),
    ("10.1. O presente Anexo \u00e9 parte integrante e indissoci\u00e1vel do Acordo de Coopera\u00e7\u00e3o Internacional celebrado entre ASAS DE SOCORRO e ADI, aplicando-se-lhe, no que couber, todas as suas disposi\u00e7\u00f5es, em especial as cl\u00e1usulas relativas \u00e0 lei aplic\u00e1vel e foro (Cl\u00e1usula 14\u00aa), confidencialidade (Cl\u00e1usula 9\u00aa) e indeniza\u00e7\u00e3o (Cl\u00e1usula 13\u00aa).",
     "10.1. This Annex is an integral and inseverable part of the International Cooperation Agreement entered into between ASAS DE SOCORRO and ADI, and all provisions thereof shall apply hereto as applicable, in particular the clauses regarding governing law and jurisdiction (Clause 14), confidentiality (Clause 9), and indemnification (Clause 13)."),
    ("An\u00e1polis/GO, _____ de __________________ de 2026.", "An\u00e1polis/GO, _______ of __________________, 2026."),
    ("Nome: ______________________________", "Name: ______________________________"),
    ("Cargo: ______________________________", "Title: ______________________________"),
    ("TESTEMUNHAS:", "WITNESSES:"),
    ("CPF: ___________________", "ID No.: ___________________"),
    ("Assinatura: ______________________________", "Signature: ______________________________"),
]


def translate_text(text):
    for pt, en in TRANSLATIONS:
        if pt in text:
            text = text.replace(pt, en)
    return text


def translate_paragraph(para):
    full_text = "".join(run.text for run in para.runs)
    translated = translate_text(full_text)
    if translated == full_text:
        return
    if para.runs:
        para.runs[0].text = translated
        for run in para.runs[1:]:
            run.text = ""


def translate_table(table):
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                translate_paragraph(para)
            for nested in cell.tables:
                translate_table(nested)


doc = Document(SRC)
for para in doc.paragraphs:
    translate_paragraph(para)
for table in doc.tables:
    translate_table(table)
doc.save(DST)
print("Saved:", DST)
